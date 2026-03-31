import streamlit as st
import pandas as pd
import plotly.graph_objects as go
from datetime import datetime
import io
import base64
from docx import Document
from docx.shared import Inches

# Importaciones de tus módulos de base de datos
from ..database.sql_db import execute_query
from ..database.semantic_mongo_db import get_student_semantic_analysis
from ..utils.helpers import decode_professor_id

def generate_comparative_report(student_name, df_student):
    """Genera un archivo Word con el progreso, comparando el primer hito con el último."""
    doc = Document()
    doc.add_heading(f'Reporte de Evolución Semántica: {student_name}', 0)
    
    # Ordenar por fecha para asegurar cronología
    df_sorted = df_student.sort_values('Fecha')
    inicio = df_sorted.iloc[0]
    final = df_sorted.iloc[-1]
    
    doc.add_heading('1. Resumen de Progreso', level=1)
    p = doc.add_paragraph()
    p.add_run(f"Estudiante: ").bold = True
    p.add_run(f"{student_name}\n")
    p.add_run(f"Periodo de Análisis: ").bold = True
    p.add_run(f"{inicio['Fecha'].strftime('%d/%m/%Y')} al {final['Fecha'].strftime('%d/%m/%Y')}\n")
    p.add_run(f"Mejora en Coherencia (M1): ").bold = True
    p.add_run(f"{((final['M1'] - inicio['M1']) * 100):+.1f}%")

    # Tabla de métricas
    doc.add_heading('2. Comparativa de Hitos Key', level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Hito'
    hdr_cells[1].text = 'Fecha/Hora'
    hdr_cells[2].text = 'M1 (Coherencia)'
    hdr_cells[3].text = 'M2 (Robustez)'

    for label, hito in [("Inicial", inicio), ("Actual", final)]:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = hito['Fecha'].strftime('%d/%m/%Y %H:%M')
        row_cells[2].text = f"{hito['M1']:.4f}"
        row_cells[3].text = f"{hito['M2']:.4f}"

    # Inserción de Imágenes de Grafos
    doc.add_heading('3. Mapas Mentales Comparativos', level=1)
    
    for label, hito in [("Mapa Semántico Inicial", inicio), ("Mapa Semántico Actual", final)]:
        doc.add_heading(f"{label} ({hito['Fecha'].strftime('%H:%M')})", level=2)
        
        # Intentar extraer la imagen guardada en el análisis
        # Nota: 'concept_graph' suele ser el diccionario, 'img_b64' es el campo de la imagen
        img_str = hito.get('img_b64')
        if img_str:
            try:
                img_data = base64.b64decode(img_str)
                img_stream = io.BytesIO(img_data)
                doc.add_picture(img_stream, width=Inches(5))
            except Exception:
                doc.add_paragraph("[Error al decodificar la imagen del grafo]")
        else:
            doc.add_paragraph("[Imagen no disponible en el registro histórico]")

    doc.add_page_break()
    doc.add_heading('4. Notas del Tutor', level=1)
    doc.add_paragraph("__________________________________________________________")
    doc.add_paragraph("Firma del Docente / Sello Institucional")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def professor_page():
    # 1. Identificación y Contexto
    prof_id = st.session_state.get('username', 'MG-EDU-UNIFE-LIM')
    full_name = st.session_state.get('full_name', 'Docente')
    info = decode_professor_id(prof_id)
    
    st.sidebar.markdown(f"### 👩‍🏫 {full_name}")
    st.sidebar.info(f"**Institución:** {info['institucion']}\n\n**Facultad:** {info['facultad']}")

    st.title("Panel del Instructor - Seguimiento Semántico")
    
    # 2. Selector de Grupos (Combo Box Dinámico)
    filtro_prefix = f"{info['iniciales']}-%"
    query_grupos = "SELECT DISTINCT class_id FROM users WHERE id LIKE %s AND role = 'Estudiante'"
    grupos_db = execute_query(query_grupos, (filtro_prefix,))
    lista_grupos = [g['class_id'] for g in grupos_db] if grupos_db else [prof_id]
    
    grupo_seleccionado = st.selectbox("Seleccione el grupo de tesistas:", lista_grupos)

    # 3. Obtención de Datos de MongoDB
    raw_data = get_student_semantic_analysis(group_id=grupo_seleccionado)
    
    if not raw_data:
        st.warning(f"No se encontraron datos semánticos para el grupo {grupo_seleccionado}.")
        return

    # Procesar datos para DataFrame
    processed_list = []
    for a in raw_data:
        res = a.get('analysis_result', {})
        # En tu estructura, M1 está en analysis_result y M2 dentro del concept_graph
        graph = res.get('concept_graph', {})
        processed_list.append({
            'Estudiante': a.get('username'),
            'Fecha': a.get('timestamp'),
            'M1': res.get('m1_score', 0),
            'M2': graph.get('M2_density', 0) if isinstance(graph, dict) else 0,
            'Fase': res.get('fase', 'N/A'),
            'Texto': a.get('text', ''),
            'img_b64': res.get('graph_image') # Campo donde se guarda el b64
        })
    
    df = pd.DataFrame(processed_list).sort_values('Fecha')

    # --- TABS: VISIÓN GENERAL vs DETALLADA ---
    tab_gen, tab_det = st.tabs(["📊 Visión General del Grupo", "🔍 Seguimiento por Alumno"])

    with tab_gen:
        st.subheader("Estado Global de la Clase")
        
        # Métricas de resumen
        c1, c2, c3 = st.columns(3)
        c1.metric("Alumnos Activos", len(df['Estudiante'].unique()))
        c2.metric("Promedio M1 (Coherencia)", f"{df['M1'].mean():.2f}")
        c3.metric("Promedio M2 (Robustez)", f"{df['M2'].mean():.2f}")

        # Gráfica Dual M1/M2 Comparativa
        fig = go.Figure()
        for student in df['Estudiante'].unique():
            sdf = df[df['Estudiante'] == student]
            fig.add_trace(go.Scatter(x=sdf['Fecha'], y=sdf['M1'], name=f"M1: {student}", mode='lines+markers'))
            fig.add_trace(go.Scatter(x=sdf['Fecha'], y=sdf['M2'], name=f"M2: {student}", mode='lines', line=dict(dash='dot'), visible='legendonly'))
        
        fig.update_layout(title="Evolución Semántica del Grupo", yaxis_title="Score (0 a 1)", hovermode="x unified")
        st.plotly_chart(fig, use_container_width=True)

    with tab_det:
        st.subheader("Expediente Detallado de Tesista")
        est_sel = st.selectbox("Seleccione Alumno:", df['Estudiante'].unique())
        
        if est_sel:
            student_df = df[df['Estudiante'] == est_sel].sort_values('Fecha')
            
            # Botón de Descarga Reporte DOCX
            report_buffer = generate_comparative_report(est_sel, student_df)
            st.download_button(
                label="📥 Descargar Reporte Evolutivo (.docx)",
                data=report_buffer,
                file_name=f"Progreso_{est_sel}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            st.divider()
            
            # Slider para navegar por las sesiones del alumno
            sesion_idx = st.select_slider(
                "Navegar historial de sesiones:",
                options=range(len(student_df)),
                format_func=lambda i: student_df.iloc[i]['Fecha'].strftime('%d/%m/%Y %H:%M')
            )
            
            sel_row = student_df.iloc[sesion_idx]
            
            col_text, col_metrics = st.columns([2, 1])
            with col_text:
                st.markdown(f"**Fase Actual:** {sel_row['Fase']}")
                st.info(f"**Texto de la sesión:**\n\n{sel_row['Texto']}")
            
            with col_metrics:
                st.metric("Coherencia M1", f"{sel_row['M1']:.4f}")
                st.metric("Robustez M2", f"{sel_row['M2']:.4f}")
                if sel_row['M1'] < 0.6:
                    st.error("Alerta: Baja Coherencia")
                elif sel_row['M1'] > 0.8:
                    st.success("Estado: Alta Alineación")