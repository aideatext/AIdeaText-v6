# modules/ui/professor/professor_ui.py
"""
Dashboard del Profesor-Asesor — Piloto UNIFE 2025

Arquitectura de 3 pestañas (Doc3_Metodo_v2):
  Tab 1: Alcance de la Investigación  — resumen cohorte + actividad
  Tab 2: M1 Coherencia Dialógica      — heatmap + timeline + drill-down
  Tab 3: M2 Robustez Estructural      — trayectoria + export + acciones

Flujo de navegación: Cohorte → Grupo → Estudiante → Sesión
"""

import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
import pandas as pd
import numpy as np
from datetime import datetime, timedelta, timezone
import logging
import io

from modules.database.sql_db import execute_query
from modules.database.semantic_mongo_db import get_student_semantic_analysis
from modules.database.semantic_mongo_live_db import get_student_semantic_live_analysis
from modules.database.discourse_mongo_db import get_student_discourse_analysis
from modules.database.chat_mongo_db import get_chat_history
from modules.metrics.metrics_processor import process_semantic_data
from modules.metrics.report_generator import create_docx_report
from modules.utils.helpers import decode_professor_id

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
#  CONSTANTES: Paleta de señales pedagógicas (azul oscuro → blanco)
# ─────────────────────────────────────────────────────────────────────────────

SIGNAL_COLORS = {
    'progreso':       '#1B2A4A',  # Azul oscuro — óptimo
    'regresion':      '#3A5A8C',  # Azul medio-oscuro
    'brecha':         '#6B8FC2',  # Azul medio
    'estancamiento':  '#A8C4E0',  # Azul claro
    'inactividad':    '#FFFFFF',  # Blanco — no debería existir
}

SIGNAL_LABELS = {
    'progreso':      'Progreso',
    'regresion':     'Regresión',
    'brecha':        'Brecha escritura-diálogo',
    'estancamiento': 'Estancamiento',
    'inactividad':   'Inactividad',
}

SIGNAL_ACTIONS = {
    'progreso':      'Exigir mayor profundidad: pedir al estudiante que desarrolle sub-argumentos en su próxima versión.',
    'regresion':     'Reforzar conceptos clave en sesión presencial o remota. Revisar si el estudiante cambió de tema o perdió el hilo.',
    'brecha':        'El estudiante dialoga conceptos que no escribe (o escribe sin integrar el diálogo). Sesión de acompañamiento en escritura.',
    'estancamiento': 'Cambiar estrategia: proponer nuevas fuentes, reformular la pregunta de investigación, o sesión de brainstorming con el tutor virtual.',
    'inactividad':   'Contactar al estudiante. Puede necesitar motivación, apoyo logístico o reorientación.',
}


# ─────────────────────────────────────────────────────────────────────────────
#  HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _parse_username(username: str) -> dict:
    """MG-G1-2025-1-t1 → {id_estudiante, id_grupo, display_grupo}"""
    if not username:
        return {'id_estudiante': 'S/N', 'id_grupo': '', 'display_grupo': ''}
    parts = username.split('-')
    if len(parts) >= 5 and parts[-1].lower().startswith('t'):
        return {
            'id_estudiante': parts[-1],
            'id_grupo': '-'.join(parts[:-1]),
            'display_grupo': f"Grupo {parts[1]} ({parts[2]}-{parts[3]})",
        }
    return {'id_estudiante': username, 'id_grupo': username, 'display_grupo': username}


def _classify_signal(m1_values: list, last_activity_date, m1_escritura: float = 0, m1_chat: float = 0) -> str:
    """
    Clasifica al estudiante en una de las 5 señales pedagógicas.
    Prioridad: inactividad > brecha > regresión > estancamiento > progreso.
    """
    now = datetime.now(timezone.utc)

    # Inactividad: sin sesiones en >7 días
    if last_activity_date:
        if hasattr(last_activity_date, 'tzinfo') and last_activity_date.tzinfo is None:
            last_activity_date = last_activity_date.replace(tzinfo=timezone.utc)
        days_inactive = (now - last_activity_date).days
        if days_inactive > 7:
            return 'inactividad'
    else:
        return 'inactividad'

    # Brecha escritura-diálogo: diferencia > 0.25 entre escritura y chat
    if m1_escritura > 0 and m1_chat > 0:
        brecha = abs(m1_escritura - m1_chat)
        if brecha > 0.25:
            return 'brecha'

    # Necesitamos al menos 2 valores de M1 para evaluar tendencia
    active = [v for v in m1_values if v > 0]
    if len(active) < 2:
        if active and active[0] >= 0.60:
            return 'progreso'
        return 'estancamiento'

    # Tendencia: comparar primera mitad vs segunda mitad
    mid = len(active) // 2
    first_half = np.mean(active[:mid])
    second_half = np.mean(active[mid:])
    delta = second_half - first_half

    if delta > 0.05:
        return 'progreso'
    elif delta < -0.05:
        return 'regresion'
    else:
        return 'estancamiento'


# ─────────────────────────────────────────────────────────────────────────────
#  DATA BUILDERS (4 cubetas → DataFrames)
# ─────────────────────────────────────────────────────────────────────────────

def _build_chat_df(group_id: str) -> pd.DataFrame:
    try:
        raw = get_chat_history(group_id=group_id, analysis_type='chat_interaction')
        if not raw:
            return pd.DataFrame()
        rows = []
        for chat in raw:
            meta = chat.get('metadata', {}) or {}
            username = chat.get('username', f"{group_id}-t?")
            p = _parse_username(username)
            rows.append({
                'Estudiante_ID':     p['id_estudiante'],
                'Username_Completo': username,
                'Grupo_ID':          p['id_grupo'],
                'Grupo_Display':     p['display_grupo'],
                'Fecha':             chat.get('timestamp', datetime.utcnow()),
                'M1_chat':           float(meta.get('m1_score', 0.0)),
                'M2_chat':           float(meta.get('m2_score', 0.0)),
                'depth_chat':        int(meta.get('depth', 0)),
                'visual_graph':      chat.get('visual_graph'),
                'graph_user':        meta.get('graph_user'),
                'graph_tutor':       meta.get('graph_tutor'),
                'n_messages':        len(chat.get('messages', [])),
                'Fuente':            'Chat',
            })
        return pd.DataFrame(rows).sort_values('Fecha')
    except Exception as e:
        logger.error(f"Error en _build_chat_df({group_id}): {e}")
        return pd.DataFrame()


def _build_live_df(group_id: str) -> pd.DataFrame:
    try:
        raw = get_student_semantic_live_analysis(group_id=group_id, limit=100)
        if not raw:
            return pd.DataFrame()
        rows = []
        for doc in raw:
            username = doc.get('username', '')
            p = _parse_username(username)
            res = doc.get('analysis_result', {}) or {}
            rows.append({
                'Estudiante_ID':     p['id_estudiante'],
                'Username_Completo': username,
                'Grupo_ID':          p['id_grupo'],
                'Grupo_Display':     p['display_grupo'],
                'Fecha':             doc.get('timestamp', datetime.utcnow()),
                'M1':                float(res.get('m1_score', 0.0)),
                'M2':                float(doc.get('m2_score', 0.0)),
                'Texto':             (doc.get('text') or '')[:300],
                'concept_graph':     doc.get('concept_graph'),
                'Tipo':              'Live',
            })
        return pd.DataFrame(rows).sort_values('Fecha') if rows else pd.DataFrame()
    except Exception as e:
        logger.error(f"Error en _build_live_df({group_id}): {e}")
        return pd.DataFrame()


def _build_discourse_df(group_id: str) -> pd.DataFrame:
    try:
        raw = get_student_discourse_analysis(group_id=group_id, limit=100)
        if not raw:
            return pd.DataFrame()
        rows = []
        for doc in raw:
            username = doc.get('username', '')
            p = _parse_username(username)
            rows.append({
                'Estudiante_ID':     p['id_estudiante'],
                'Username_Completo': username,
                'Grupo_ID':          p['id_grupo'],
                'Grupo_Display':     p['display_grupo'],
                'Fecha':             doc.get('timestamp', datetime.utcnow()),
                'Texto1':            (doc.get('text1') or '')[:200],
                'Texto2':            (doc.get('text2') or '')[:200],
                'key_concepts1':     doc.get('key_concepts1', []),
                'key_concepts2':     doc.get('key_concepts2', []),
                'graph1':            doc.get('graph1'),
                'graph2':            doc.get('graph2'),
                'Tipo':              'Comparado',
            })
        return pd.DataFrame(rows).sort_values('Fecha') if rows else pd.DataFrame()
    except Exception as e:
        logger.error(f"Error en _build_discourse_df({group_id}): {e}")
        return pd.DataFrame()


# ─────────────────────────────────────────────────────────────────────────────
#  PÁGINA PRINCIPAL
# ─────────────────────────────────────────────────────────────────────────────

def professor_page():
    prof_user = st.session_state.get('username', 'MG-EDU-UNIFE-LIM')
    full_name = st.session_state.get('full_name', 'Docente')

    try:
        info = decode_professor_id(prof_user)
    except Exception:
        info = {'institucion': 'UNIFE', 'facultad': 'Educación', 'iniciales': 'MG'}

    # ── Sidebar ──
    with st.sidebar:
        st.markdown(f"### {full_name}")
        st.divider()
        st.markdown(f"**Institución:** {info['institucion']}")
        st.markdown(f"**Facultad:** {info['facultad']}")
        st.divider()
        st.caption(f"Cátedra: {info['iniciales']}")

    st.title("Panel de Gestión Académica")

    # ── Grupos (SQL) ──
    user_prefix = info['iniciales']
    query = """
        SELECT DISTINCT class_id FROM users
        WHERE (id LIKE %s OR class_id LIKE %s)
        AND role = 'Estudiante'
        AND class_id != %s
    """
    grupos_db = execute_query(query, (f"{user_prefix}%", f"{user_prefix}%", prof_user))
    if not grupos_db:
        st.warning("No se encontraron grupos de tesistas bajo su cátedra.")
        return
    lista_grupos_ids = sorted(list(set(g['class_id'] for g in grupos_db if g['class_id'])))
    if not lista_grupos_ids:
        st.warning("No se encontraron grupos de tesistas bajo su cátedra.")
        return

    # ── Carga de las 4 cubetas ──
    all_raw_sem = []
    chat_frames, live_frames, disc_frames = [], [], []

    for g_id in lista_grupos_ids:
        all_raw_sem.extend(get_student_semantic_analysis(group_id=g_id))
        df_c = _build_chat_df(g_id)
        if not df_c.empty: chat_frames.append(df_c)
        df_l = _build_live_df(g_id)
        if not df_l.empty: live_frames.append(df_l)
        df_d = _build_discourse_df(g_id)
        if not df_d.empty: disc_frames.append(df_d)

    df_sem      = process_semantic_data(all_raw_sem)
    df_chat_all = pd.concat(chat_frames, ignore_index=True) if chat_frames else pd.DataFrame()
    df_live_all = pd.concat(live_frames, ignore_index=True) if live_frames else pd.DataFrame()
    df_disc_all = pd.concat(disc_frames, ignore_index=True) if disc_frames else pd.DataFrame()

    # ── 3 Pestañas ──
    tab_alcance, tab_m1, tab_m2 = st.tabs([
        "Alcance de la Investigación",
        "M1 — Coherencia Dialógica",
        "M2 — Robustez Estructural",
    ])

    with tab_alcance:
        _render_tab_alcance(df_sem, df_chat_all, df_live_all, df_disc_all, lista_grupos_ids, info)

    with tab_m1:
        _render_tab_m1(df_sem, df_chat_all, df_live_all, df_disc_all)

    with tab_m2:
        _render_tab_m2(df_sem, df_chat_all, df_live_all, lista_grupos_ids)


# ═════════════════════════════════════════════════════════════════════════════
#  TAB 1: ALCANCE DE LA INVESTIGACIÓN
# ═════════════════════════════════════════════════════════════════════════════

def _render_tab_alcance(df_sem, df_chat, df_live, df_disc, grupos, info):
    st.subheader("Alcance de la Investigación")
    st.caption(f"Piloto {info['institucion']} — Cohorte Abril–Junio 2025")

    # Contadores de actividad
    n_estudiantes = len(set(
        list(df_sem['Username_Completo'].unique() if not df_sem.empty else []) +
        list(df_chat['Username_Completo'].unique() if not df_chat.empty else []) +
        list(df_live['Username_Completo'].unique() if not df_live.empty else [])
    ))

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Grupos", len(grupos))
    c2.metric("Estudiantes", n_estudiantes)
    c3.metric("Análisis Semánticos", len(df_sem))
    c4.metric("Análisis Live", len(df_live))
    c5.metric("Checkpoints Tutor", len(df_chat))

    st.divider()

    # Tabla de actividad por grupo
    st.markdown("**Actividad por grupo**")
    grupo_data = []
    for g_id in grupos:
        p = _parse_username(g_id + '-t0')
        n_sem = len(df_sem[df_sem['Grupo_ID'] == g_id]) if not df_sem.empty and 'Grupo_ID' in df_sem.columns else 0
        n_live = len(df_live[df_live['Grupo_ID'] == g_id]) if not df_live.empty and 'Grupo_ID' in df_live.columns else 0
        n_chat = len(df_chat[df_chat['Grupo_ID'] == g_id]) if not df_chat.empty and 'Grupo_ID' in df_chat.columns else 0
        n_disc = len(df_disc[df_disc['Grupo_ID'] == g_id]) if not df_disc.empty and 'Grupo_ID' in df_disc.columns else 0

        # Estudiantes en grupo
        est_en_grupo = set()
        for df in [df_sem, df_live, df_chat]:
            if not df.empty and 'Username_Completo' in df.columns and 'Grupo_ID' in df.columns:
                est_en_grupo.update(df[df['Grupo_ID'] == g_id]['Username_Completo'].unique())

        # Última actividad
        fechas = []
        for df in [df_sem, df_live, df_chat]:
            if not df.empty and 'Grupo_ID' in df.columns and 'Fecha' in df.columns:
                fdf = df[df['Grupo_ID'] == g_id]
                if not fdf.empty:
                    fechas.append(fdf['Fecha'].max())
        ultima = max(fechas) if fechas else None

        # M1 promedio del grupo
        m1_avg = 0.0
        if not df_sem.empty and 'Grupo_ID' in df_sem.columns:
            gdf = df_sem[df_sem['Grupo_ID'] == g_id]
            if not gdf.empty and 'M1' in gdf.columns:
                m1_avg = gdf['M1'].mean()

        grupo_data.append({
            'Grupo': p['display_grupo'],
            'Estudiantes': len(est_en_grupo),
            'Semánticos': n_sem,
            'Live': n_live,
            'Tutor': n_chat,
            'Comparados': n_disc,
            'M1 Prom.': round(m1_avg, 3),
            'Última actividad': ultima.strftime('%d/%m/%Y') if ultima and hasattr(ultima, 'strftime') else '—',
        })

    if grupo_data:
        st.dataframe(pd.DataFrame(grupo_data), use_container_width=True, hide_index=True)

    # Semáforo general de señales pedagógicas
    st.divider()
    st.markdown("**Señales pedagógicas — Vista general**")
    _render_signal_summary(df_sem, df_chat, df_live)


def _render_signal_summary(df_sem, df_chat, df_live):
    """Clasifica a cada estudiante y muestra el conteo de señales."""
    all_users = set()
    for df in [df_sem, df_chat, df_live]:
        if not df.empty and 'Username_Completo' in df.columns:
            all_users.update(df['Username_Completo'].unique())

    if not all_users:
        st.info("Sin datos suficientes para clasificar señales.")
        return

    signal_counts = {k: 0 for k in SIGNAL_COLORS}

    for user in all_users:
        m1_vals = []
        last_date = None

        # M1 de escritura
        m1_esc = 0.0
        if not df_sem.empty and 'Username_Completo' in df_sem.columns:
            udf = df_sem[df_sem['Username_Completo'] == user]
            if not udf.empty:
                m1_vals.extend(udf['M1'].tolist())
                m1_esc = udf['M1'].mean()
                last_date = udf['Fecha'].max()

        # M1 de chat
        m1_ch = 0.0
        if not df_chat.empty and 'Username_Completo' in df_chat.columns:
            uchat = df_chat[df_chat['Username_Completo'] == user]
            if not uchat.empty:
                m1_vals.extend(uchat['M1_chat'].tolist())
                m1_ch = uchat['M1_chat'].mean()
                chat_date = uchat['Fecha'].max()
                if last_date is None or chat_date > last_date:
                    last_date = chat_date

        signal = _classify_signal(m1_vals, last_date, m1_esc, m1_ch)
        signal_counts[signal] += 1

    # Mostrar barras horizontales con colores
    cols = st.columns(len(SIGNAL_COLORS))
    for i, (key, color) in enumerate(SIGNAL_COLORS.items()):
        count = signal_counts[key]
        text_color = '#FFFFFF' if key in ('progreso', 'regresion') else '#1B2A4A'
        border = f"border: 1px solid #1B2A4A;" if key == 'inactividad' else ""
        with cols[i]:
            st.markdown(
                f"""<div style='background:{color};{border}border-radius:8px;
                padding:12px;text-align:center;min-height:90px;'>
                <p style='color:{text_color};font-size:2em;font-weight:bold;margin:0;'>{count}</p>
                <p style='color:{text_color};font-size:0.8em;margin:4px 0 0 0;'>{SIGNAL_LABELS[key]}</p>
                </div>""",
                unsafe_allow_html=True,
            )


# ═════════════════════════════════════════════════════════════════════════════
#  TAB 2: M1 — COHERENCIA DIALÓGICA
# ═════════════════════════════════════════════════════════════════════════════

def _render_tab_m1(df_sem, df_chat, df_live, df_disc):
    st.subheader("M1 — Coherencia Dialógica")
    st.caption(
        "Mide la coherencia entre lo que el estudiante escribe y su interacción "
        "con el tutor virtual. Calculado con similitud coseno (modelo multilingüe 768 dims)."
    )

    if df_sem.empty and df_chat.empty and df_live.empty:
        st.info("Sin datos de M1 todavía. Los estudiantes deben realizar análisis y usar el tutor virtual.")
        return

    # ── Unificar todos los M1 en un solo timeline ──
    timeline_rows = []

    if not df_sem.empty:
        for _, r in df_sem.iterrows():
            timeline_rows.append({
                'Estudiante_ID': r.get('Estudiante_ID', ''),
                'Username': r.get('Username_Completo', ''),
                'Grupo': r.get('Grupo_Display', ''),
                'Grupo_ID': r.get('Grupo_ID', ''),
                'Fecha': r.get('Fecha'),
                'M1': r.get('M1', 0.0),
                'Tipo': 'Semántico',
            })

    if not df_live.empty and 'M1' in df_live.columns:
        for _, r in df_live.iterrows():
            if r.get('M1', 0) > 0:
                timeline_rows.append({
                    'Estudiante_ID': r.get('Estudiante_ID', ''),
                    'Username': r.get('Username_Completo', ''),
                    'Grupo': r.get('Grupo_Display', ''),
                    'Grupo_ID': r.get('Grupo_ID', ''),
                    'Fecha': r.get('Fecha'),
                    'M1': r.get('M1', 0.0),
                    'Tipo': 'Live',
                })

    if not df_chat.empty:
        for _, r in df_chat.iterrows():
            if r.get('M1_chat', 0) > 0:
                timeline_rows.append({
                    'Estudiante_ID': r.get('Estudiante_ID', ''),
                    'Username': r.get('Username_Completo', ''),
                    'Grupo': r.get('Grupo_Display', ''),
                    'Grupo_ID': r.get('Grupo_ID', ''),
                    'Fecha': r.get('Fecha'),
                    'M1': r.get('M1_chat', 0.0),
                    'Tipo': 'Tutor Virtual',
                })

    if not timeline_rows:
        st.info("Sin valores M1 registrados todavía.")
        return

    df_timeline = pd.DataFrame(timeline_rows).sort_values('Fecha')

    # ── NIVEL COHORTE: Heatmap temporal ──
    st.markdown("### Vista Cohorte — Mapa de calor")

    # Preparar datos para heatmap: eje X = semana, eje Y = estudiante
    df_hm = df_timeline.copy()
    df_hm['Semana'] = pd.to_datetime(df_hm['Fecha']).dt.isocalendar().week
    heatmap_data = df_hm.groupby(['Estudiante_ID', 'Semana'])['M1'].mean().reset_index()

    if not heatmap_data.empty:
        hm_pivot = heatmap_data.pivot(index='Estudiante_ID', columns='Semana', values='M1')
        hm_pivot = hm_pivot.fillna(np.nan)

        fig_hm = go.Figure(data=go.Heatmap(
            z=hm_pivot.values,
            x=[f"Sem {int(c)}" for c in hm_pivot.columns],
            y=hm_pivot.index.tolist(),
            colorscale=[
                [0.0, '#FFFFFF'],
                [0.2, '#A8C4E0'],
                [0.4, '#6B8FC2'],
                [0.6, '#3A5A8C'],
                [0.8, '#1B2A4A'],
                [1.0, '#0D1B2A'],
            ],
            zmin=0, zmax=1,
            colorbar=dict(title='M1'),
            hoverongaps=False,
            hovertemplate='%{y}<br>%{x}<br>M1: %{z:.3f}<extra></extra>',
        ))
        fig_hm.update_layout(
            title='Coherencia Dialógica por Estudiante y Semana',
            xaxis_title='Semana del semestre',
            yaxis_title='Estudiante',
            height=max(300, len(hm_pivot) * 40 + 100),
        )
        st.plotly_chart(fig_hm, use_container_width=True)

    st.divider()

    # ── NIVEL GRUPO: Timeline por grupo ──
    st.markdown("### Vista Grupo — Evolución M1")

    grupos_disponibles = sorted(df_timeline['Grupo'].unique())
    if not grupos_disponibles:
        st.info("Sin datos de grupo.")
        return

    grupo_sel = st.selectbox("Seleccionar grupo:", grupos_disponibles, key="m1_grupo_sel")
    df_grupo = df_timeline[df_timeline['Grupo'] == grupo_sel]

    if not df_grupo.empty:
        # Símbolos por tipo de interacción
        symbol_map = {'Semántico': 'circle', 'Live': 'diamond', 'Tutor Virtual': 'square'}

        fig_grupo = px.scatter(
            df_grupo, x='Fecha', y='M1',
            color='Estudiante_ID',
            symbol='Tipo',
            symbol_map=symbol_map,
            title=f"Evolución M1 — {grupo_sel}",
            labels={'M1': 'M1 Coherencia Dialógica', 'Fecha': ''},
        )
        # Líneas conectando puntos por estudiante
        for est in df_grupo['Estudiante_ID'].unique():
            df_est = df_grupo[df_grupo['Estudiante_ID'] == est].sort_values('Fecha')
            fig_grupo.add_trace(go.Scatter(
                x=df_est['Fecha'], y=df_est['M1'],
                mode='lines', line=dict(width=1, dash='dot'),
                showlegend=False, opacity=0.4,
            ))

        fig_grupo.update_layout(
            yaxis=dict(range=[0, 1.05]),
            legend=dict(orientation='h', y=-0.15),
            height=450,
        )
        # Líneas de referencia
        fig_grupo.add_hline(y=0.80, line_dash="dash", line_color="#1B2A4A",
                            annotation_text="Alta (0.80)", annotation_position="top left")
        fig_grupo.add_hline(y=0.60, line_dash="dash", line_color="#6B8FC2",
                            annotation_text="Moderada-alta (0.60)", annotation_position="bottom left")
        st.plotly_chart(fig_grupo, use_container_width=True)

    st.divider()

    # ── NIVEL ESTUDIANTE: Detalle con grafos ──
    st.markdown("### Vista Estudiante — Detalle por sesión")

    estudiantes_grupo = sorted(df_grupo['Username'].unique()) if not df_grupo.empty else []
    if not estudiantes_grupo:
        st.info("Sin estudiantes con datos M1 en este grupo.")
        return

    est_sel = st.selectbox(
        "Seleccionar estudiante:",
        estudiantes_grupo,
        format_func=lambda u: f"{_parse_username(u)['id_estudiante']} ({u})",
        key="m1_est_sel",
    )

    df_est = df_grupo[df_grupo['Username'] == est_sel].sort_values('Fecha')

    if df_est.empty:
        st.info("Sin sesiones M1 para este estudiante.")
        return

    # Tabla de sesiones
    st.dataframe(
        df_est[['Fecha', 'Tipo', 'M1']].assign(
            Fecha=df_est['Fecha'].apply(
                lambda x: x.strftime('%d/%m/%Y %H:%M') if hasattr(x, 'strftime') else str(x)
            ),
            M1=df_est['M1'].apply(lambda x: f"{x:.4f}"),
        ),
        use_container_width=True,
        hide_index=True,
    )

    # Selector de sesión para ver grafo
    _render_student_session_detail(est_sel, df_sem, df_chat, df_live, df_disc)


def _render_student_session_detail(username, df_sem, df_chat, df_live, df_disc):
    """Muestra detalle de sesión seleccionada: tipo, fecha, M1, grafo."""

    # Sub-tabs por tipo de interacción
    sub_sem, sub_live, sub_comp, sub_chat = st.tabs([
        "Análisis Semántico",
        "Análisis en Vivo",
        "Análisis Comparado",
        "Interacción Tutor Virtual",
    ])

    est_id = _parse_username(username)['id_estudiante']

    # [1] Semántico
    with sub_sem:
        df_es = df_sem[df_sem['Username_Completo'] == username] if not df_sem.empty else pd.DataFrame()
        if df_es.empty:
            st.info("Sin análisis semánticos registrados.")
        else:
            df_es = df_es.sort_values('Fecha')
            idx = st.select_slider(
                "Sesión:", options=range(len(df_es)),
                format_func=lambda i: (
                    df_es.iloc[i]['Fecha'].strftime('%d/%m %H:%M')
                    if hasattr(df_es.iloc[i]['Fecha'], 'strftime')
                    else str(df_es.iloc[i]['Fecha'])
                ),
                key="m1_sem_slider",
            )
            sesion = df_es.iloc[idx]
            col_t, col_m = st.columns([2, 1])
            with col_t:
                st.markdown("**Texto analizado:**")
                st.info(sesion.get('Texto', ''))
            with col_m:
                st.metric("M1 Coherencia", f"{sesion.get('M1', 0):.4f}")
                st.metric("M2 Densidad", f"{sesion.get('M2', 0):.4f}")

    # [2] Live
    with sub_live:
        df_el = df_live[df_live['Username_Completo'] == username] if not df_live.empty else pd.DataFrame()
        if df_el.empty:
            st.info("Sin análisis live registrados.")
        else:
            df_el = df_el.sort_values('Fecha')
            idx_l = st.select_slider(
                "Sesión live:", options=range(len(df_el)),
                format_func=lambda i: (
                    df_el.iloc[i]['Fecha'].strftime('%d/%m %H:%M')
                    if hasattr(df_el.iloc[i]['Fecha'], 'strftime')
                    else str(df_el.iloc[i]['Fecha'])
                ),
                key="m1_live_slider",
            )
            sesion_l = df_el.iloc[idx_l]
            col_tl, col_ml = st.columns([2, 1])
            with col_tl:
                st.markdown("**Texto live:**")
                st.info(sesion_l.get('Texto', ''))
                if sesion_l.get('concept_graph'):
                    st.image(sesion_l['concept_graph'], caption="Grafo Live", use_container_width=True)
            with col_ml:
                st.metric("M2 Densidad", f"{sesion_l.get('M2', 0):.4f}")

    # [3] Comparado
    with sub_comp:
        df_ed = df_disc[df_disc['Username_Completo'] == username] if not df_disc.empty else pd.DataFrame()
        if df_ed.empty:
            st.info("Sin análisis comparados registrados.")
        else:
            df_ed = df_ed.sort_values('Fecha')
            idx_d = st.select_slider(
                "Sesión comparada:", options=range(len(df_ed)),
                format_func=lambda i: (
                    df_ed.iloc[i]['Fecha'].strftime('%d/%m %H:%M')
                    if hasattr(df_ed.iloc[i]['Fecha'], 'strftime')
                    else str(df_ed.iloc[i]['Fecha'])
                ),
                key="m1_disc_slider",
            )
            sesion_d = df_ed.iloc[idx_d]
            col_d1, col_d2 = st.columns(2)
            with col_d1:
                st.markdown("**Documento 1:**")
                st.info(sesion_d.get('Texto1', ''))
                if sesion_d.get('graph1'):
                    st.image(sesion_d['graph1'], caption="Grafo Doc 1", use_container_width=True)
            with col_d2:
                st.markdown("**Documento 2:**")
                st.info(sesion_d.get('Texto2', ''))
                if sesion_d.get('graph2'):
                    st.image(sesion_d['graph2'], caption="Grafo Doc 2", use_container_width=True)

    # [4] Tutor Virtual
    with sub_chat:
        df_ec = df_chat[df_chat['Username_Completo'] == username] if not df_chat.empty else pd.DataFrame()
        if df_ec.empty:
            st.info("Sin checkpoints de tutor registrados.")
        else:
            df_ec = df_ec.sort_values('Fecha')
            idx_c = st.select_slider(
                "Checkpoint:", options=range(len(df_ec)),
                format_func=lambda i: (
                    df_ec.iloc[i]['Fecha'].strftime('%d/%m %H:%M')
                    if hasattr(df_ec.iloc[i]['Fecha'], 'strftime')
                    else str(df_ec.iloc[i]['Fecha'])
                ),
                key="m1_chat_slider",
            )
            sesion_c = df_ec.iloc[idx_c]

            cm1, cm2 = st.columns(2)
            cm1.metric("M1 Coherencia Dialógica", f"{sesion_c.get('M1_chat', 0):.4f}")
            cm2.metric("M2 Densidad Chat", f"{sesion_c.get('M2_chat', 0):.4f}")
            st.caption(f"Turnos en el checkpoint: {sesion_c.get('n_messages', '?')}")

            st.markdown("---")

            gc1, gc2, gc3 = st.columns(3)
            with gc1:
                st.markdown("**Grafo del Estudiante**")
                if sesion_c.get('graph_user'):
                    st.image(sesion_c['graph_user'], caption="G_u", use_container_width=True)
                else:
                    st.caption("No disponible")
            with gc2:
                st.markdown("**Grafo Compartido (Sincronía)**")
                if sesion_c.get('visual_graph'):
                    st.image(sesion_c['visual_graph'], caption="G_hybrid", use_container_width=True)
                else:
                    st.caption("No disponible")
            with gc3:
                st.markdown("**Grafo del Tutor**")
                if sesion_c.get('graph_tutor'):
                    st.image(sesion_c['graph_tutor'], caption="G_t", use_container_width=True)
                else:
                    st.caption("No disponible")


# ═════════════════════════════════════════════════════════════════════════════
#  TAB 3: M2 — ROBUSTEZ ESTRUCTURAL
# ═════════════════════════════════════════════════════════════════════════════

def _render_tab_m2(df_sem, df_chat, df_live, grupos):
    st.subheader("M2 — Robustez Estructural")
    st.caption(
        "Mide el crecimiento en densidad, alineación temática y profundidad argumentativa "
        "del grafo del estudiante a lo largo del semestre."
    )

    if df_sem.empty:
        st.info("Sin datos de análisis semántico. M2 se calcula sobre los grafos de escritura.")
        return

    # ── NIVEL GRUPO: M2 promedio ──
    st.markdown("### Vista por Grupo")

    grupos_disponibles = sorted(df_sem['Grupo_Display'].unique())
    grupo_m2_sel = st.selectbox("Seleccionar grupo:", grupos_disponibles, key="m2_grupo_sel")
    df_g = df_sem[df_sem['Grupo_Display'] == grupo_m2_sel]

    if df_g.empty:
        st.info("Sin datos para este grupo.")
        return

    # Tabla resumen por estudiante
    resumen_est = df_g.groupby('Estudiante_ID').agg(
        Sesiones=('M1', 'count'),
        M1_prom=('M1', 'mean'),
        M2_prom=('M2', 'mean'),
        Ultima_fecha=('Fecha', 'max'),
    ).reset_index()
    resumen_est['M1_prom'] = resumen_est['M1_prom'].round(4)
    resumen_est['M2_prom'] = resumen_est['M2_prom'].round(4)
    resumen_est['Ultima_fecha'] = resumen_est['Ultima_fecha'].apply(
        lambda x: x.strftime('%d/%m/%Y') if hasattr(x, 'strftime') else str(x)
    )
    resumen_est.columns = ['Estudiante', 'Sesiones', 'M1 Promedio', 'M2 Densidad Prom.', 'Última sesión']

    st.dataframe(resumen_est, use_container_width=True, hide_index=True)

    # Gráfico M2 por estudiante
    fig_m2_bar = px.bar(
        resumen_est, x='Estudiante', y='M2 Densidad Prom.',
        title=f"M2 Densidad Promedio — {grupo_m2_sel}",
        color='M2 Densidad Prom.',
        color_continuous_scale=[[0, '#FFFFFF'], [0.5, '#6B8FC2'], [1, '#1B2A4A']],
    )
    fig_m2_bar.update_layout(yaxis=dict(range=[0, max(0.5, resumen_est['M2 Densidad Prom.'].max() * 1.2)]))
    st.plotly_chart(fig_m2_bar, use_container_width=True)

    # Exportar informe de grupo
    try:
        group_report_buf = _generate_group_report(grupo_m2_sel, df_g, df_chat)
        st.download_button(
            label=f"Descargar informe — {grupo_m2_sel}",
            data=group_report_buf,
            file_name=f"Informe_{grupo_m2_sel.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="m2_group_report",
        )
    except Exception as e:
        logger.error(f"Error generando informe de grupo: {e}")

    st.divider()

    # ── NIVEL ESTUDIANTE: Trayectoria M2 ──
    st.markdown("### Vista por Estudiante — Trayectoria M2")

    estudiantes_grupo = sorted(df_g['Username_Completo'].unique())
    est_m2_sel = st.selectbox(
        "Seleccionar estudiante:",
        estudiantes_grupo,
        format_func=lambda u: f"{_parse_username(u)['id_estudiante']} ({u})",
        key="m2_est_sel",
    )

    df_est = df_g[df_g['Username_Completo'] == est_m2_sel].sort_values('Fecha')

    if not df_est.empty:
        # Gráfico de trayectoria M1 + M2
        fig_traj = go.Figure()
        fig_traj.add_trace(go.Scatter(
            x=df_est['Fecha'], y=df_est['M1'],
            mode='lines+markers', name='M1 Coherencia',
            line=dict(color='#1B2A4A', width=2),
            marker=dict(size=8),
        ))
        fig_traj.add_trace(go.Scatter(
            x=df_est['Fecha'], y=df_est['M2'],
            mode='lines+markers', name='M2 Densidad',
            line=dict(color='#6B8FC2', width=2),
            marker=dict(size=8),
        ))
        fig_traj.update_layout(
            title=f"Trayectoria M1/M2 — {_parse_username(est_m2_sel)['id_estudiante']}",
            yaxis=dict(range=[0, 1.05], title='Score'),
            xaxis_title='',
            legend=dict(orientation='h', y=-0.15),
            height=400,
        )
        st.plotly_chart(fig_traj, use_container_width=True)

        # Exportar informe individual
        try:
            est_report_buf = create_docx_report(est_m2_sel, df_est)
            st.download_button(
                label=f"Descargar informe — {_parse_username(est_m2_sel)['id_estudiante']}",
                data=est_report_buf,
                file_name=f"Informe_{est_m2_sel}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="m2_est_report",
            )
        except Exception as e:
            logger.error(f"Error generando informe individual: {e}")

    st.divider()

    # ── Panel de Acción Pedagógica ──
    _render_action_panel(df_g, df_chat, grupo_m2_sel)


def _render_action_panel(df_grupo, df_chat, grupo_display):
    """
    Panel de indicaciones de acción para el profesor.
    Clasifica cada estudiante del grupo y sugiere acciones concretas.
    """
    st.markdown(f"### Indicaciones de Acción — {grupo_display}")

    all_users = sorted(df_grupo['Username_Completo'].unique())
    if not all_users:
        return

    for user in all_users:
        p = _parse_username(user)
        udf = df_grupo[df_grupo['Username_Completo'] == user]

        m1_vals = udf['M1'].tolist()
        last_date = udf['Fecha'].max() if not udf.empty else None

        m1_esc = udf['M1'].mean() if not udf.empty else 0.0
        m1_ch = 0.0
        if not df_chat.empty and 'Username_Completo' in df_chat.columns:
            uchat = df_chat[df_chat['Username_Completo'] == user]
            if not uchat.empty:
                m1_ch = uchat['M1_chat'].mean()

        signal = _classify_signal(m1_vals, last_date, m1_esc, m1_ch)
        color = SIGNAL_COLORS[signal]
        label = SIGNAL_LABELS[signal]
        action = SIGNAL_ACTIONS[signal]
        text_color = '#FFFFFF' if signal in ('progreso', 'regresion') else '#1B2A4A'
        border = "border: 1px solid #1B2A4A;" if signal == 'inactividad' else ""

        m1_display = f"{m1_esc:.3f}" if m1_esc > 0 else "—"

        st.markdown(
            f"""<div style='background:{color};{border}border-radius:8px;
            padding:14px;margin-bottom:10px;'>
            <div style='display:flex;justify-content:space-between;align-items:center;'>
                <div>
                    <span style='color:{text_color};font-weight:bold;font-size:1.1em;'>
                        {p['id_estudiante']}
                    </span>
                    <span style='color:{text_color};font-size:0.9em;margin-left:12px;'>
                        M1: {m1_display} | Sesiones: {len(m1_vals)}
                    </span>
                </div>
                <span style='color:{text_color};font-weight:bold;'>{label}</span>
            </div>
            <p style='color:{text_color};font-size:0.85em;margin:6px 0 0 0;
            opacity:0.9;'>{action}</p>
            </div>""",
            unsafe_allow_html=True,
        )


# ─────────────────────────────────────────────────────────────────────────────
#  GENERADOR DE INFORME DE GRUPO
# ─────────────────────────────────────────────────────────────────────────────

def _generate_group_report(grupo_display, df_grupo, df_chat):
    """Genera un DOCX con el resumen del grupo."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(f'Informe de Grupo: {grupo_display}', 0)
    doc.add_paragraph(f'Generado: {datetime.now().strftime("%d/%m/%Y %H:%M")}')

    # Resumen
    doc.add_heading('Resumen', level=1)
    n_est = df_grupo['Estudiante_ID'].nunique()
    n_ses = len(df_grupo)
    m1_avg = df_grupo['M1'].mean() if not df_grupo.empty else 0
    m2_avg = df_grupo['M2'].mean() if not df_grupo.empty and 'M2' in df_grupo.columns else 0

    doc.add_paragraph(f'Estudiantes: {n_est}')
    doc.add_paragraph(f'Total sesiones: {n_ses}')
    doc.add_paragraph(f'M1 Coherencia promedio: {m1_avg:.4f}')
    doc.add_paragraph(f'M2 Densidad promedio: {m2_avg:.4f}')

    # Tabla por estudiante
    doc.add_heading('Detalle por Estudiante', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Estudiante'
    hdr[1].text = 'Sesiones'
    hdr[2].text = 'M1 Prom.'
    hdr[3].text = 'M2 Prom.'

    for est_id in sorted(df_grupo['Estudiante_ID'].unique()):
        edf = df_grupo[df_grupo['Estudiante_ID'] == est_id]
        row = table.add_row().cells
        row[0].text = est_id
        row[1].text = str(len(edf))
        row[2].text = f"{edf['M1'].mean():.4f}"
        row[3].text = f"{edf['M2'].mean():.4f}" if 'M2' in edf.columns else "—"

    # Señales pedagógicas
    doc.add_heading('Señales Pedagógicas', level=1)
    for user in sorted(df_grupo['Username_Completo'].unique()):
        p = _parse_username(user)
        udf = df_grupo[df_grupo['Username_Completo'] == user]
        m1_vals = udf['M1'].tolist()
        last_date = udf['Fecha'].max() if not udf.empty else None
        m1_esc = udf['M1'].mean() if not udf.empty else 0.0
        m1_ch = 0.0
        if not df_chat.empty and 'Username_Completo' in df_chat.columns:
            uchat = df_chat[df_chat['Username_Completo'] == user]
            if not uchat.empty:
                m1_ch = uchat['M1_chat'].mean()

        signal = _classify_signal(m1_vals, last_date, m1_esc, m1_ch)
        doc.add_paragraph(
            f"{p['id_estudiante']} — {SIGNAL_LABELS[signal]}: {SIGNAL_ACTIONS[signal]}"
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
