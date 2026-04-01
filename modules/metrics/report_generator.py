#modules/metrics/report_generator.py
from docx import Document
from docx.shared import Inches
import io
import base64

def create_docx_report(student_name, df_student):
    doc = Document()
    doc.add_heading(f'Evolución Semántica: {student_name}', 0)
    
    # Tomamos el primero y el último cronológicamente
    inicio = df_student.iloc[0]
    final = df_student.iloc[-1]
    
    doc.add_heading('Comparativa de Hitos', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Métrica'
    hdr[1].text = 'Inicial (' + inicio['Fecha'].strftime('%d/%m') + ')'
    hdr[2].text = 'Actual (' + final['Fecha'].strftime('%d/%m') + ')'

    for m in [('M1 (Coherencia)', 'M1'), ('M2 (Robustez)', 'M2')]:
        row = table.add_row().cells
        row[0].text = m[0]
        row[1].text = f"{inicio[m[1]]:.4f}"
        row[2].text = f"{final[m[1]]:.4f}"

    doc.add_heading('3. Mapas Semánticos del progreso de la tesis', level=1)
        
    for label, hito in [("Estado Inicial", df_student.iloc[0]), ("Estado Actual", df_student.iloc[-1])]:
            doc.add_heading(label, level=2)
            img_data = hito.get('img_raw')
            
            if img_data:
                try:
                    # Limpieza de Base64
                    if isinstance(img_data, str):
                        if "base64," in img_data:
                            img_data = img_data.split("base64,")[1]
                        img_data = base64.b64decode(img_data.strip())
                    
                    # Inserción en Word
                    image_stream = io.BytesIO(img_data)
                    doc.add_picture(image_stream, width=Inches(4.5))
                except Exception as e:
                    doc.add_paragraph(f"[No se pudo renderizar la imagen: {str(e)}]")
            else:
                doc.add_paragraph("[Sin evidencia visual registrada]")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer