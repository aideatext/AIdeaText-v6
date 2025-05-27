import streamlit as st
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import io

def export_data(user_data, t, format='pdf'):
    if format == 'pdf':
        return export_to_pdf(user_data, t)
    elif format == 'docx':
        return export_to_docx(user_data, t)
    else:
        raise ValueError(f"Unsupported format: {format}")

def export_to_pdf(user_data, t):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    # Título
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, t['analysis_report'])

    # Resumen
    c.setFont("Helvetica", 12)
    c.drawString(50, height - 80, f"{t['morpho_analyses']}: {len(user_data['morphosyntax_analyses'])}")
    c.drawString(50, height - 100, f"{t['semantic_analyses']}: {len(user_data['semantic_analyses'])}")
    c.drawString(50, height - 120, f"{t['discourse_analyses']}: {len(user_data['discourse_analyses'])}")

    # Aquí agregarías más detalles de los análisis...

    c.save()
    buffer.seek(0)
    return buffer

def export_to_docx(user_data, t):
    doc = Document()
    doc.add_heading(t['analysis_report'], 0)

    doc.add_paragraph(f"{t['morpho_analyses']}: {len(user_data['morphosyntax_analyses'])}")
    doc.add_paragraph(f"{t['semantic_analyses']}: {len(user_data['semantic_analyses'])}")
    doc.add_paragraph(f"{t['discourse_analyses']}: {len(user_data['discourse_analyses'])}")

    # Aquí agregarías más detalles de los análisis...

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def display_export_options(t):
    format = st.radio(t['select_export_format'], ['PDF', 'DOCX'])
    if st.button(t['export']):
        user_data = st.session_state.user_data
        if format == 'PDF':
            buffer = export_data(user_data, t, format='pdf')
            st.download_button(
                label=t['download_pdf'],
                data=buffer,
                file_name="analysis_report.pdf",
                mime="application/pdf"
            )
        elif format == 'DOCX':
            buffer = export_data(user_data, t, format='docx')
            st.download_button(
                label=t['download_docx'],
                data=buffer,
                file_name="analysis_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )